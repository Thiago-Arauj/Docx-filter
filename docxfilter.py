import os
import docx


def filter_files(searches):


    path = r"YOUR_PATH_HERE"
    files = os.listdir(path)



    for file in files:
        # THIS MATCH CASE IS FOR IN CASE YOUR FOLDER HAS ANOTHER FOLDERS IN IT, SO IN EACH CASE INSERT THE NAME OF EACH FOLDER  
        match file:
            case FOLDER1:
                path_alt1 = f"{path}/2022"
                files_alt1 = os.listdir(path_alt1)
                
                # ON CONDITION INSERT WHAT WORD YOU'D LIKE TO FILTER, FOR EXAMPLE "TEST"
                if (CONDITION) and ".docx" in file:

                    for file_alt in files_alt1:
                        with open(path_alt1+f"/{file_alt}", "r",encoding="utf-8") as doc:
                            
                                doc_read = docx.Document(path_alt1+f"/{file_alt}")

                                for para in doc_read.paragraphs:

                                    for search in searches:
                                        if search in para.text:
                                            print(file_alt)

                
                print("Pasta antiga 2022")



            case FOLDER2:
                path_alt2 = f"{path}/2023"
                files_alt2 = os.listdir(path_alt2)

                # ON CONDITION INSERT WHAT WORD YOU'D LIKE TO FILTER, FOR EXAMPLE "TEST"
                if (CONDITION) and ".docx" in file:

                    for file_alt in files_alt2:
                        if "Prova" in file_alt or "Avaliação" in file_alt or "Avaliacao" in file_alt:

                            with open(path_alt2+f"/{file_alt}", "r",encoding="utf-8") as doc:
                            
                                doc_read = docx.Document(path_alt2+f"/{file_alt}")

                                for para in doc_read.paragraphs:

                                    for search in searches:
                                        if search in para.text:
                                            print(file_alt)

                print("Pasta antiga 2023")
                


            # IF THE FILE DOES NOT HAVE ANY FOLDERS IN IT USE JUST THIS BLOCK
            case _:
                # ON CONDITION INSERT WHAT WORD YOU'D LIKE TO FILTER, FOR EXAMPLE "TEST"
                if (CONDITION) and ".docx" in file:
                    
                    with open(path+f"/{file}", "r",encoding="utf-8") as doc:
                        
                        doc_read = docx.Document(path+f"/{file}")

                        for para in doc_read.paragraphs:

                            for search in searches:
                                if search in para.text:
                                    print(file)



filter_files(searches=[YOUR,SEARCHES, HERE])
